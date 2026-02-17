"""Multi-format text extraction service for document ingestion."""
import io
import logging
from pypdf import PdfReader
from docx import Document
from bs4 import BeautifulSoup
import html2text

logger = logging.getLogger(__name__)


def normalize_text(text: str) -> str:
    """
    Normalize extracted text to handle legal document notation.

    Handles:
    - Superscript numbers (e.g., 509⁶ → "509⁶ (509-6)" - keeps original + adds searchable version)
    - Preserves original formatting while adding normalized versions
    """
    if not text:
        return text

    import re

    # Superscript to regular number mapping
    superscript_map = {
        '⁰': '0', '¹': '1', '²': '2', '³': '3', '⁴': '4',
        '⁵': '5', '⁶': '6', '⁷': '7', '⁸': '8', '⁹': '9'
    }

    # Match number followed by one or more superscripts
    # Example: "497²⁶" → base_num="497", superscripts="²⁶"
    pattern = r'(\d+)([⁰¹²³⁴⁵⁶⁷⁸⁹]+)'

    def replace_superscript(match):
        base_num = match.group(1)
        superscripts = match.group(2)
        # Convert superscripts to regular numbers: "²⁶" → "26"
        regular_nums = ''.join(superscript_map.get(s, s) for s in superscripts)
        # Keep original and add normalized version: "497²⁶ (497-26)"
        return f"{base_num}{superscripts} ({base_num}-{regular_nums})"

    # Run replacement once on entire text
    normalized = re.sub(pattern, replace_superscript, text)
    return normalized


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from PDF using pypdf.

    Args:
        file_bytes: Raw PDF file bytes

    Returns:
        Extracted text content

    Raises:
        ValueError: If PDF extraction fails or no text found
    """
    try:
        pdf_reader = PdfReader(io.BytesIO(file_bytes))

        if len(pdf_reader.pages) == 0:
            raise ValueError("PDF has no pages")

        text_parts = []
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            except Exception as e:
                logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                continue

        if not text_parts:
            raise ValueError("No text content extracted from PDF (might be scanned images or empty)")

        full_text = "\n\n".join(text_parts)

        if not full_text.strip():
            raise ValueError("PDF appears to be empty or contains only images")

        return normalize_text(full_text)

    except ValueError:
        raise
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from DOCX using python-docx.

    Args:
        file_bytes: Raw DOCX file bytes

    Returns:
        Extracted text content

    Raises:
        ValueError: If DOCX extraction fails or no text found
    """
    try:
        doc = Document(io.BytesIO(file_bytes))

        text_parts = []

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        if not text_parts:
            raise ValueError("No text content extracted from DOCX (document appears to be empty)")

        full_text = "\n\n".join(text_parts)

        if not full_text.strip():
            raise ValueError("DOCX appears to be empty")

        return normalize_text(full_text)

    except ValueError:
        raise
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")


def extract_text_from_html(file_bytes: bytes) -> str:
    """
    Extract text from HTML using BeautifulSoup and html2text.

    Converts HTML to Markdown format to preserve structure (headings, lists, etc.).

    Args:
        file_bytes: Raw HTML file bytes

    Returns:
        Extracted text content (Markdown format)

    Raises:
        ValueError: If HTML extraction fails or no text found
    """
    try:
        # Decode HTML with error handling
        try:
            html_content = file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            html_content = file_bytes.decode('utf-8', errors='ignore')
            logger.warning("HTML file had encoding issues, used ignore mode")

        # Convert HTML to Markdown (preserves structure)
        h = html2text.HTML2Text()
        h.ignore_links = False  # Keep links
        h.ignore_images = True  # Skip images
        h.ignore_emphasis = False  # Keep bold/italic
        h.body_width = 0  # No line wrapping

        markdown_text = h.handle(html_content)

        if not markdown_text.strip():
            # Fallback: Try plain text extraction with BeautifulSoup
            soup = BeautifulSoup(html_content, 'html.parser')

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.extract()

            plain_text = soup.get_text(separator='\n\n', strip=True)

            if not plain_text.strip():
                raise ValueError("No text content extracted from HTML (document appears to be empty)")

            return normalize_text(plain_text)

        return normalize_text(markdown_text)

    except ValueError:
        raise
    except Exception as e:
        logger.error(f"HTML extraction failed: {e}")
        raise ValueError(f"Failed to extract text from HTML: {str(e)}")


def extract_text(file_bytes: bytes, file_type: str) -> str:
    """
    Extract text from file bytes based on file type.

    Main dispatcher function that routes to format-specific extractors.

    Args:
        file_bytes: Raw file bytes
        file_type: MIME type of the file

    Returns:
        Extracted text content

    Raises:
        ValueError: If extraction fails or file type is unsupported

    Supported formats:
        - text/plain (.txt)
        - text/markdown (.md)
        - application/pdf (.pdf)
        - application/vnd.openxmlformats-officedocument.wordprocessingml.document (.docx)
        - text/html (.html, .htm)
    """
    logger.debug(f"Extracting text from file type: {file_type}")

    try:
        # Plain text and Markdown (existing support)
        if file_type in ("text/plain", "text/markdown"):
            try:
                text = file_bytes.decode("utf-8")
            except UnicodeDecodeError:
                text = file_bytes.decode("utf-8", errors="ignore")
                logger.warning("Text file had encoding issues, used ignore mode")

            if not text.strip():
                raise ValueError("File is empty")

            return text

        # PDF (new support)
        elif file_type == "application/pdf":
            return extract_text_from_pdf(file_bytes)

        # DOCX (new support)
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return extract_text_from_docx(file_bytes)

        # HTML (new support)
        elif file_type == "text/html":
            return extract_text_from_html(file_bytes)

        # Unsupported type
        else:
            raise ValueError(
                f"Unsupported file type: {file_type}. "
                f"Supported types: .txt, .md, .pdf, .docx, .html"
            )

    except ValueError:
        raise
    except Exception as e:
        logger.error(f"Text extraction failed for {file_type}: {e}")
        raise ValueError(f"Failed to extract text: {str(e)}")
